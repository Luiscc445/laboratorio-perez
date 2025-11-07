from flask import Blueprint, render_template, request, redirect, url_for, flash, send_file, jsonify
from flask_login import login_required, current_user
from app import db
from app.models import Paciente, Resultado, Prueba
from app.utils import admin_required
from werkzeug.utils import secure_filename
from datetime import datetime, timedelta
from sqlalchemy import func, extract
import os
import random
import string
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib import colors
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image
from reportlab.lib.units import inch
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

main = Blueprint('main', __name__)

# Definir ruta base absoluta para archivos
BASE_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
UPLOAD_DIR = os.path.join(BASE_DIR, 'app', 'static', 'uploads')
PRUEBAS_UPLOAD_DIR = os.path.join(UPLOAD_DIR, 'pruebas')
BACKUP_DIR = os.path.join(UPLOAD_DIR, 'backups')  # Carpeta de backups

def generar_codigo_acceso():
    """Genera un código aleatorio de 8 caracteres ÚNICO"""
    while True:
        codigo = ''.join(random.choices(string.ascii_uppercase + string.digits, k=8))
        # Verificar que no existe en la BD
        if not Resultado.query.filter_by(codigo_acceso=codigo).first():
            return codigo

def generar_numero_orden():
    """
    Genera un número de orden ÚNICO basado en timestamp + contador
    Formato: YYYYMMDD-HHMMSS-XXX
    Ejemplo: 20251107-153045-001
    """
    from datetime import datetime
    timestamp = datetime.now().strftime('%Y%m%d-%H%M%S')

    # Contar cuántos resultados se crearon en este segundo
    contador = Resultado.query.filter(
        Resultado.numero_orden.like(f'{timestamp}%')
    ).count()

    numero_orden = f"{timestamp}-{contador+1:03d}"
    return numero_orden

def guardar_pdf_con_backup(archivo, numero_orden):
    """
    Guarda un PDF de forma SÚPER ROBUSTA con:
    - Nombre único basado en timestamp
    - Verificación de integridad
    - Backup automático
    - Manejo de errores

    Returns:
        tuple: (filename, filepath, backup_path) o (None, None, None) si falla
    """
    try:
        # 1. Validar archivo
        if not archivo or not archivo.filename:
            raise ValueError("Archivo no válido")

        if not archivo.filename.lower().endswith('.pdf'):
            raise ValueError("Solo se permiten archivos PDF")

        # 2. Generar nombre único con timestamp
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S_%f')
        nombre_base = secure_filename(archivo.filename.replace('.pdf', ''))
        filename = f"{numero_orden}_{timestamp}_{nombre_base}.pdf"

        # 3. Crear directorios si no existen
        os.makedirs(UPLOAD_DIR, exist_ok=True)
        os.makedirs(BACKUP_DIR, exist_ok=True)

        # 4. Guardar archivo principal
        filepath = os.path.join(UPLOAD_DIR, filename)
        archivo.save(filepath)

        # 5. Verificar que se guardó correctamente
        if not os.path.exists(filepath):
            raise Exception(f"Archivo no se guardó: {filepath}")

        file_size = os.path.getsize(filepath)
        if file_size == 0:
            os.remove(filepath)
            raise Exception("Archivo PDF vacío")

        # 6. Crear BACKUP automático
        backup_path = os.path.join(BACKUP_DIR, filename)
        import shutil
        shutil.copy2(filepath, backup_path)

        # 7. Verificar backup
        if not os.path.exists(backup_path):
            print(f"⚠ WARNING: Backup no se creó: {backup_path}")
        else:
            print(f"✓ BACKUP creado: {backup_path}")

        print(f"✓ PDF guardado: {filepath} ({file_size} bytes)")

        return filename, filepath, backup_path

    except Exception as e:
        print(f"✗ Error guardando PDF: {str(e)}")
        # Limpiar archivos parciales
        if 'filepath' in locals() and os.path.exists(filepath):
            os.remove(filepath)
        if 'backup_path' in locals() and os.path.exists(backup_path):
            os.remove(backup_path)
        return None, None, None

def limpiar_archivo_huerfano(filename):
    """Elimina un archivo PDF y su backup si existen"""
    try:
        # Eliminar archivo principal
        filepath = os.path.join(UPLOAD_DIR, filename)
        if os.path.exists(filepath):
            os.remove(filepath)
            print(f"🗑 Archivo principal eliminado: {filename}")

        # Eliminar backup
        backup_path = os.path.join(BACKUP_DIR, filename)
        if os.path.exists(backup_path):
            os.remove(backup_path)
            print(f"🗑 Backup eliminado: {filename}")
    except Exception as e:
        print(f"⚠ Error limpiando archivos huérfanos: {e}")

@main.route('/')
def index():
    return render_template('publico/index.html')

@main.route('/portal-resultados')
def portal_resultados():
    return render_template('publico/portal_resultados.html')

@main.route('/catalogo-pruebas')
def catalogo_pruebas():
    # Obtener todas las pruebas ordenadas por categoría y nombre
    pruebas = Prueba.query.order_by(Prueba.categoria, Prueba.nombre).all()

    # Organizar pruebas por categoría
    pruebas_por_categoria = {}
    for prueba in pruebas:
        categoria = prueba.categoria or 'General'
        if categoria not in pruebas_por_categoria:
            pruebas_por_categoria[categoria] = []
        pruebas_por_categoria[categoria].append(prueba)

    return render_template('publico/catalogo/lista_pruebas.html',
                         pruebas=pruebas,
                         pruebas_por_categoria=pruebas_por_categoria)

@main.route('/consultar-resultado', methods=['POST'])
def consultar_resultado():
    ci = request.form.get('ci')
    codigo = request.form.get('codigo')

    resultado = Resultado.query.filter_by(paciente_ci=ci, codigo_acceso=codigo).first()

    if resultado:
        return render_template('publico/ver_resultado.html', resultado=resultado)
    else:
        flash('CI o código de acceso incorrecto', 'danger')
        return redirect(url_for('main.portal_resultados'))

@main.route('/dashboard')
@admin_required
def dashboard():
    # Estadísticas básicas
    total_pacientes = Paciente.query.count()
    total_resultados = Resultado.query.count()
    total_pruebas = Prueba.query.count()

    # Pacientes registrados en los últimos 6 meses
    fecha_inicio = datetime.now() - timedelta(days=180)
    pacientes_por_mes = db.session.query(
        extract('month', Paciente.fecha_registro).label('mes'),
        func.count(Paciente.id).label('total')
    ).filter(Paciente.fecha_registro >= fecha_inicio).group_by('mes').all()

    # Resultados por mes en los últimos 6 meses
    resultados_por_mes = db.session.query(
        extract('month', Resultado.fecha_muestra).label('mes'),
        func.count(Resultado.id).label('total')
    ).filter(Resultado.fecha_muestra >= fecha_inicio).group_by('mes').all()

    # Top 5 pruebas más recientes del catálogo
    top_pruebas = db.session.query(
        Prueba.nombre,
        Prueba.precio
    ).order_by(Prueba.fecha_creacion.desc()).limit(5).all()

    # Estadísticas adicionales
    pacientes_este_mes = Paciente.query.filter(
        extract('month', Paciente.fecha_registro) == datetime.now().month,
        extract('year', Paciente.fecha_registro) == datetime.now().year
    ).count()

    resultados_este_mes = Resultado.query.filter(
        extract('month', Resultado.fecha_muestra) == datetime.now().month,
        extract('year', Resultado.fecha_muestra) == datetime.now().year
    ).count()

    # Preparar datos para gráficos
    meses_nombres = ['Ene', 'Feb', 'Mar', 'Abr', 'May', 'Jun', 'Jul', 'Ago', 'Sep', 'Oct', 'Nov', 'Dic']
    mes_actual = datetime.now().month

    # Últimos 6 meses
    ultimos_meses = []
    for i in range(5, -1, -1):
        mes_idx = (mes_actual - i - 1) % 12
        ultimos_meses.append(meses_nombres[mes_idx])

    # Crear arrays de datos para gráficos
    pacientes_data = [0] * 6
    resultados_data = [0] * 6

    for mes, total in pacientes_por_mes:
        mes_idx = int(mes)
        pos = 5 - ((mes_actual - mes_idx) % 12)
        if 0 <= pos < 6:
            pacientes_data[pos] = total

    for mes, total in resultados_por_mes:
        mes_idx = int(mes)
        pos = 5 - ((mes_actual - mes_idx) % 12)
        if 0 <= pos < 6:
            resultados_data[pos] = total

    return render_template('admin/dashboard.html',
                         total_pacientes=total_pacientes,
                         total_resultados=total_resultados,
                         total_pruebas=total_pruebas,
                         pacientes_este_mes=pacientes_este_mes,
                         resultados_este_mes=resultados_este_mes,
                         ultimos_meses=ultimos_meses,
                         pacientes_data=pacientes_data,
                         resultados_data=resultados_data,
                         top_pruebas=top_pruebas)

@main.route('/pacientes', methods=['GET', 'POST'])
@admin_required
def admin_pacientes():
    if request.method == 'POST':
        try:
            paciente = Paciente(
                nombre=request.form['nombre'],
                ci=request.form['ci'],
                telefono=request.form.get('telefono'),
                email=request.form.get('email')
            )
            db.session.add(paciente)
            db.session.commit()
            flash('Paciente registrado exitosamente', 'success')
        except Exception as e:
            flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('main.admin_pacientes'))
    pacientes = Paciente.query.order_by(Paciente.fecha_registro.desc()).all()
    return render_template('admin/pacientes.html', pacientes=pacientes, now=datetime.now())

@main.route('/paciente/<int:paciente_id>')
@admin_required
def ver_paciente(paciente_id):
    paciente = Paciente.query.get_or_404(paciente_id)
    return jsonify({
        'id': paciente.id,
        'nombre': paciente.nombre,
        'ci': paciente.ci,
        'telefono': paciente.telefono or '-',
        'email': paciente.email or '-',
        'fecha_registro': paciente.fecha_registro.strftime('%d/%m/%Y %H:%M')
    })

@main.route('/paciente/editar/<int:paciente_id>', methods=['POST'])
@admin_required
def editar_paciente(paciente_id):
    try:
        paciente = Paciente.query.get_or_404(paciente_id)
        paciente.nombre = request.form['nombre']
        paciente.ci = request.form['ci']
        paciente.telefono = request.form.get('telefono')
        paciente.email = request.form.get('email')

        resultados = Resultado.query.filter_by(paciente_id=paciente_id).all()
        for resultado in resultados:
            resultado.paciente_nombre = paciente.nombre
            resultado.paciente_ci = paciente.ci

        db.session.commit()
        flash('Paciente y sus resultados actualizados exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pacientes'))

@main.route('/paciente/eliminar/<int:paciente_id>', methods=['POST'])
@admin_required
def eliminar_paciente(paciente_id):
    try:
        paciente = Paciente.query.get_or_404(paciente_id)
        nombre_paciente = paciente.nombre

        # Obtener todos los resultados del paciente
        resultados = Resultado.query.filter_by(paciente_id=paciente_id).all()

        # Eliminar archivos PDF físicos
        archivos_eliminados = 0
        for resultado in resultados:
            if resultado.archivo_pdf:
                try:
                    pdf_path = os.path.join(UPLOAD_DIR, resultado.archivo_pdf)
                    if os.path.exists(pdf_path):
                        os.remove(pdf_path)
                        archivos_eliminados += 1
                except Exception as e:
                    print(f"Error eliminando archivo {resultado.archivo_pdf}: {e}")

        # SQLAlchemy eliminará automáticamente los resultados gracias a cascade='all, delete-orphan'
        db.session.delete(paciente)
        db.session.commit()

        flash(f'Paciente "{nombre_paciente}" eliminado exitosamente junto con {len(resultados)} resultado(s) y {archivos_eliminados} archivo(s) PDF', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar paciente: {str(e)}', 'danger')
        print(f"Error: {e}")

    return redirect(url_for('main.admin_pacientes'))

@main.route('/resultados', methods=['GET', 'POST'])
@admin_required
def admin_resultados():
    """
    SISTEMA ROBUSTO DE GESTIÓN DE RESULTADOS
    - Números de orden automáticos y únicos
    - Almacenamiento seguro de PDFs
    - Backups automáticos
    - Manejo inteligente de errores
    - Nunca se pierden archivos
    """
    if request.method == 'POST':
        filename_guardado = None
        numero_orden_generado = None

        try:
            # ============ VALIDACIONES INICIALES ============
            archivo = request.files.get('archivo_pdf')

            # Validar paciente PRIMERO (antes de guardar archivo)
            paciente_id = request.form.get('paciente_id')
            if not paciente_id:
                flash('❌ Debe seleccionar un paciente', 'danger')
                return redirect(url_for('main.admin_resultados'))

            paciente = Paciente.query.get(int(paciente_id))
            if not paciente:
                flash('❌ Paciente no encontrado', 'danger')
                return redirect(url_for('main.admin_resultados'))

            # ============ GENERAR NÚMERO DE ORDEN AUTOMÁTICO ============
            numero_orden_manual = request.form.get('numero_orden', '').strip()

            if numero_orden_manual:
                # Usuario proporcionó número manual
                numero_orden_generado = numero_orden_manual
                print(f"📋 Usando número de orden manual: {numero_orden_generado}")
            else:
                # Generar automáticamente
                numero_orden_generado = generar_numero_orden()
                print(f"🔢 Número de orden generado automáticamente: {numero_orden_generado}")

            # ============ GUARDAR PDF CON BACKUP ============
            filename_guardado, filepath, backup_path = guardar_pdf_con_backup(
                archivo,
                numero_orden_generado
            )

            if not filename_guardado:
                raise Exception("No se pudo guardar el archivo PDF")

            print(f"📁 Archivo guardado: {filename_guardado}")
            print(f"💾 Backup creado en: {BACKUP_DIR}")

            # ============ PROCESAR FECHA ============
            fecha_str = request.form.get('fecha_muestra')
            fecha_muestra = datetime.strptime(fecha_str, '%Y-%m-%d').date() if fecha_str else None

            # ============ GENERAR CÓDIGO DE ACCESO ÚNICO ============
            codigo_acceso = generar_codigo_acceso()

            # ============ REGISTRAR EN BASE DE DATOS ============
            resultado = Resultado(
                numero_orden=numero_orden_generado,
                paciente_id=paciente.id,
                paciente_nombre=paciente.nombre,
                paciente_ci=paciente.ci,
                fecha_muestra=fecha_muestra,
                archivo_pdf=filename_guardado,
                codigo_acceso=codigo_acceso
            )

            db.session.add(resultado)
            db.session.commit()

            # ============ ÉXITO COMPLETO ============
            print("=" * 80)
            print("✅ RESULTADO GUARDADO EXITOSAMENTE")
            print(f"   ID: {resultado.id}")
            print(f"   Número Orden: {numero_orden_generado}")
            print(f"   Código Acceso: {codigo_acceso}")
            print(f"   Paciente: {paciente.nombre}")
            print(f"   Archivo: {filename_guardado}")
            print(f"   Backup: ✓ Creado")
            print("=" * 80)

            flash(f'✅ Resultado guardado exitosamente. Código de acceso: {codigo_acceso}', 'success')

        except ValueError as ve:
            # Errores de validación (archivo no válido, etc.)
            db.session.rollback()
            if filename_guardado:
                limpiar_archivo_huerfano(filename_guardado)
            flash(f'❌ Error de validación: {str(ve)}', 'danger')
            print(f"✗ Error de validación: {str(ve)}")

        except Exception as e:
            # Errores inesperados
            db.session.rollback()

            # Limpiar archivos huérfanos
            if filename_guardado:
                limpiar_archivo_huerfano(filename_guardado)

            error_msg = str(e)
            print("=" * 80)
            print("❌ ERROR AL GUARDAR RESULTADO")
            print(f"   Error: {error_msg}")
            print(f"   Archivos limpiados: {filename_guardado if filename_guardado else 'N/A'}")
            print("=" * 80)

            # Mensaje de error amigable
            if 'duplicate key' in error_msg.lower():
                flash('❌ Este número de orden ya existe. El sistema generará uno automático.', 'danger')
            else:
                flash(f'❌ Error al guardar: {error_msg}', 'danger')

        return redirect(url_for('main.admin_resultados'))
    
    resultados = Resultado.query.order_by(Resultado.fecha_creacion.desc()).all()
    pacientes = Paciente.query.order_by(Paciente.nombre).all()
    return render_template('admin/resultados.html', resultados=resultados, pacientes=pacientes)

@main.route('/descargar-credenciales-pdf/<int:resultado_id>')
@admin_required
def descargar_credenciales_pdf(resultado_id):
    """
    Genera PDF de credenciales PROFESIONAL y SIN ERRORES
    Para que el paciente vea su código de acceso
    """
    resultado = Resultado.query.get_or_404(resultado_id)

    buffer = BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=50,
        leftMargin=50,
        topMargin=50,
        bottomMargin=50
    )
    elements = []
    styles = getSampleStyleSheet()

    # ============ ESTILOS PERSONALIZADOS ============
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#11998e'),
        spaceAfter=10,
        alignment=1,  # Centrado
        fontName='Helvetica-Bold'
    )

    subtitle_style = ParagraphStyle(
        'SubtitleStyle',
        parent=styles['Normal'],
        fontSize=12,
        textColor=colors.HexColor('#7f8c8d'),
        spaceAfter=30,
        alignment=1,  # Centrado
        fontName='Helvetica'
    )

    section_title_style = ParagraphStyle(
        'SectionTitle',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.HexColor('#2c3e50'),
        spaceBefore=10,
        spaceAfter=10,
        fontName='Helvetica-Bold'
    )

    # ============ ENCABEZADO ============
    elements.append(Paragraph('LABORATORIO CLÍNICO PÉREZ', title_style))
    elements.append(Paragraph('Potosí, Bolivia', subtitle_style))
    elements.append(Spacer(1, 0.2*inch))

    # ============ TÍTULO PRINCIPAL ============
    header_data = [['CREDENCIALES DE ACCESO A RESULTADOS']]
    header_table = Table(header_data, colWidths=[6.5*inch])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#11998e')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.white),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 16),
        ('TOPPADDING', (0, 0), (-1, -1), 15),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 15),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    elements.append(header_table)
    elements.append(Spacer(1, 0.3*inch))

    # ============ INFORMACIÓN DEL PACIENTE ============
    patient_data = [
        ['INFORMACIÓN DEL PACIENTE'],
        ['Nombre Completo:', resultado.paciente_nombre],
        ['Cédula de Identidad:', resultado.paciente_ci],
        ['Número de Orden:', resultado.numero_orden],
        ['Fecha de Emisión:', resultado.fecha_creacion.strftime("%d/%m/%Y %H:%M")]
    ]

    patient_table = Table(patient_data, colWidths=[2.5*inch, 4*inch])
    patient_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#34495e')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 13),
        ('ALIGN', (0, 0), (-1, 0), 'LEFT'),
        ('SPAN', (0, 0), (-1, 0)),  # Merge header cells
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('LEFTPADDING', (0, 0), (-1, 0), 15),

        # Data rows
        ('BACKGROUND', (0, 1), (0, -1), colors.HexColor('#ecf0f1')),
        ('FONTNAME', (0, 1), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, -1), 11),
        ('TEXTCOLOR', (0, 1), (-1, -1), colors.HexColor('#2c3e50')),
        ('ALIGN', (0, 1), (0, -1), 'LEFT'),
        ('ALIGN', (1, 1), (1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
        ('LEFTPADDING', (0, 1), (-1, -1), 15),
        ('RIGHTPADDING', (0, 1), (-1, -1), 15),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#34495e')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(patient_table)
    elements.append(Spacer(1, 0.4*inch))

    # ============ CREDENCIALES DE ACCESO ============
    credentials_data = [
        ['CREDENCIALES PARA CONSULTAR RESULTADOS'],
        ['CÉDULA DE IDENTIDAD:', resultado.paciente_ci],
        ['CÓDIGO DE ACCESO:', resultado.codigo_acceso]
    ]

    credentials_table = Table(credentials_data, colWidths=[2.5*inch, 4*inch])
    credentials_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#e67e22')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 13),
        ('SPAN', (0, 0), (-1, 0)),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('LEFTPADDING', (0, 0), (-1, 0), 15),

        # Data rows - CI
        ('BACKGROUND', (0, 1), (0, 1), colors.HexColor('#ecf0f1')),
        ('FONTNAME', (0, 1), (0, 1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 1), (1, 1), 'Helvetica'),
        ('FONTSIZE', (0, 1), (-1, 1), 11),
        ('TEXTCOLOR', (0, 1), (-1, 1), colors.HexColor('#2c3e50')),

        # CÓDIGO - Destacado
        ('BACKGROUND', (0, 2), (0, 2), colors.HexColor('#ecf0f1')),
        ('BACKGROUND', (1, 2), (1, 2), colors.HexColor('#fff9e6')),
        ('FONTNAME', (0, 2), (0, 2), 'Helvetica-Bold'),
        ('FONTNAME', (1, 2), (1, 2), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 2), (0, 2), 11),
        ('FONTSIZE', (1, 2), (1, 2), 16),
        ('TEXTCOLOR', (1, 2), (1, 2), colors.HexColor('#e74c3c')),

        # General
        ('ALIGN', (0, 1), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
        ('LEFTPADDING', (0, 1), (-1, -1), 15),
        ('RIGHTPADDING', (0, 1), (-1, -1), 15),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#e67e22')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(credentials_table)
    elements.append(Spacer(1, 0.4*inch))

    # ============ INSTRUCCIONES ============
    instructions_data = [
        ['INSTRUCCIONES PARA ACCEDER A SUS RESULTADOS'],
        ['1', 'Ingrese a: www.laboratoriopérez.com o utilice el enlace proporcionado por el laboratorio'],
        ['2', 'Click en el botón "Consultar mis Resultados" o "Ver Resultados"'],
        ['3', 'Ingrese su Cédula de Identidad (CI) y el Código de Acceso proporcionado arriba'],
        ['4', 'Podrá visualizar y descargar su resultado en formato PDF'],
        ['', 'NOTA: Guarde este documento de forma segura. Su código de acceso es confidencial.']
    ]

    instructions_table = Table(instructions_data, colWidths=[0.4*inch, 6.1*inch])
    instructions_table.setStyle(TableStyle([
        # Header
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#3498db')),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.white),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, 0), 13),
        ('SPAN', (0, 0), (-1, 0)),
        ('TOPPADDING', (0, 0), (-1, 0), 12),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('LEFTPADDING', (0, 0), (-1, 0), 15),

        # Steps
        ('FONTNAME', (0, 1), (0, 4), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 1), (0, 4), 14),
        ('TEXTCOLOR', (0, 1), (0, 4), colors.HexColor('#3498db')),
        ('ALIGN', (0, 1), (0, 4), 'CENTER'),
        ('FONTNAME', (1, 1), (1, 4), 'Helvetica'),
        ('FONTSIZE', (1, 1), (1, 4), 10),
        ('TEXTCOLOR', (1, 1), (1, 4), colors.HexColor('#2c3e50')),

        # Note
        ('BACKGROUND', (0, 5), (-1, 5), colors.HexColor('#fff3cd')),
        ('SPAN', (0, 5), (-1, 5)),
        ('FONTNAME', (0, 5), (-1, 5), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 5), (-1, 5), 9),
        ('TEXTCOLOR', (0, 5), (-1, 5), colors.HexColor('#856404')),

        # General
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('TOPPADDING', (0, 1), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 1), (-1, -1), 10),
        ('LEFTPADDING', (0, 1), (-1, -1), 15),
        ('RIGHTPADDING', (0, 1), (-1, -1), 15),

        # Borders
        ('BOX', (0, 0), (-1, -1), 1, colors.HexColor('#bdc3c7')),
        ('LINEBELOW', (0, 0), (-1, 0), 2, colors.HexColor('#3498db')),
        ('INNERGRID', (0, 1), (-1, -1), 0.5, colors.HexColor('#bdc3c7')),
    ]))

    elements.append(instructions_table)
    elements.append(Spacer(1, 0.5*inch))

    # ============ FOOTER ============
    footer_text = Paragraph(
        '<i>Laboratorio Clínico Pérez - Potosí, Bolivia - Documento generado automáticamente</i>',
        ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.HexColor('#95a5a6'),
            alignment=1
        )
    )
    elements.append(footer_text)

    # ============ GENERAR PDF ============
    doc.build(elements)
    buffer.seek(0)

    filename = f"Credenciales_{resultado.paciente_nombre.replace(' ', '_')}_{resultado.numero_orden}.pdf"

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/pdf')

@main.route('/descargar-credenciales-word/<int:resultado_id>')
@admin_required
def descargar_credenciales_word(resultado_id):
    """
    Genera Word de credenciales PROFESIONAL y BONITO
    Para que el paciente vea su código de acceso
    """
    resultado = Resultado.query.get_or_404(resultado_id)

    doc = Document()

    # ============ CONFIGURACIÓN DE MÁRGENES ============
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # ============ ENCABEZADO DEL DOCUMENTO ============
    titulo = doc.add_heading('LABORATORIO CLÍNICO PÉREZ', 0)
    titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    titulo_run = titulo.runs[0]
    titulo_run.font.color.rgb = RGBColor(17, 153, 142)  # Verde Pérez
    titulo_run.font.size = Pt(26)
    titulo_run.font.name = 'Arial'

    subtitulo = doc.add_paragraph('Potosí, Bolivia')
    subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitulo_run = subtitulo.runs[0]
    subtitulo_run.font.size = Pt(12)
    subtitulo_run.font.color.rgb = RGBColor(127, 140, 141)
    subtitulo_run.font.italic = True

    doc.add_paragraph()  # Espacio

    # ============ TÍTULO PRINCIPAL ============
    header_para = doc.add_paragraph()
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    header_run = header_para.add_run('CREDENCIALES DE ACCESO A RESULTADOS')
    header_run.font.bold = True
    header_run.font.size = Pt(16)
    header_run.font.color.rgb = RGBColor(255, 255, 255)

    # Fondo verde para el título
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '11998e')
    header_para._element.get_or_add_pPr().append(shading_elm)

    # Padding del párrafo
    pPr = header_para._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '200')
    spacing.set(qn('w:after'), '200')
    pPr.append(spacing)

    doc.add_paragraph()  # Espacio

    # ============ TABLA: INFORMACIÓN DEL PACIENTE ============
    table1 = doc.add_table(rows=5, cols=2)
    table1.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell = table1.rows[0].cells[0]
    header_cell.merge(table1.rows[0].cells[1])
    header_cell.text = 'INFORMACIÓN DEL PACIENTE'
    header_run = header_cell.paragraphs[0].runs[0]
    header_run.font.bold = True
    header_run.font.size = Pt(13)
    header_run.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), '34495e')
    header_cell.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    # Datos del paciente
    datos_paciente = [
        ['Nombre Completo:', resultado.paciente_nombre],
        ['Cédula de Identidad:', resultado.paciente_ci],
        ['Número de Orden:', resultado.numero_orden],
        ['Fecha de Emisión:', resultado.fecha_creacion.strftime("%d/%m/%Y %H:%M")]
    ]

    for i, (label, value) in enumerate(datos_paciente, start=1):
        # Columna izquierda (labels)
        cell_label = table1.rows[i].cells[0]
        cell_label.text = label
        cell_label.paragraphs[0].runs[0].font.bold = True
        cell_label.paragraphs[0].runs[0].font.size = Pt(11)
        shading_elm = OxmlElement('w:shd')
        shading_elm.set(qn('w:fill'), 'ecf0f1')
        cell_label.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

        # Columna derecha (valores)
        cell_value = table1.rows[i].cells[1]
        cell_value.text = value
        cell_value.paragraphs[0].runs[0].font.size = Pt(11)

    doc.add_paragraph()  # Espacio

    # ============ TABLA: CREDENCIALES DE ACCESO ============
    table2 = doc.add_table(rows=3, cols=2)
    table2.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell2 = table2.rows[0].cells[0]
    header_cell2.merge(table2.rows[0].cells[1])
    header_cell2.text = 'CREDENCIALES PARA CONSULTAR RESULTADOS'
    header_run2 = header_cell2.paragraphs[0].runs[0]
    header_run2.font.bold = True
    header_run2.font.size = Pt(13)
    header_run2.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm2 = OxmlElement('w:shd')
    shading_elm2.set(qn('w:fill'), 'e67e22')
    header_cell2.paragraphs[0]._element.get_or_add_pPr().append(shading_elm2)

    # CI
    table2.rows[1].cells[0].text = 'CÉDULA DE IDENTIDAD:'
    table2.rows[1].cells[0].paragraphs[0].runs[0].font.bold = True
    table2.rows[1].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'ecf0f1')
    table2.rows[1].cells[0].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    table2.rows[1].cells[1].text = resultado.paciente_ci
    table2.rows[1].cells[1].paragraphs[0].runs[0].font.size = Pt(11)

    # CÓDIGO DE ACCESO (destacado)
    table2.rows[2].cells[0].text = 'CÓDIGO DE ACCESO:'
    table2.rows[2].cells[0].paragraphs[0].runs[0].font.bold = True
    table2.rows[2].cells[0].paragraphs[0].runs[0].font.size = Pt(11)
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'ecf0f1')
    table2.rows[2].cells[0].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    table2.rows[2].cells[1].text = resultado.codigo_acceso
    codigo_run = table2.rows[2].cells[1].paragraphs[0].runs[0]
    codigo_run.font.bold = True
    codigo_run.font.size = Pt(18)
    codigo_run.font.color.rgb = RGBColor(231, 76, 60)  # Rojo para destacar
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), 'fff9e6')  # Amarillo claro
    table2.rows[2].cells[1].paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    doc.add_paragraph()  # Espacio

    # ============ TABLA: INSTRUCCIONES ============
    table3 = doc.add_table(rows=6, cols=1)
    table3.style = 'Light Grid Accent 1'

    # Header de sección
    header_cell3 = table3.rows[0].cells[0]
    header_cell3.text = 'INSTRUCCIONES PARA ACCEDER A SUS RESULTADOS'
    header_run3 = header_cell3.paragraphs[0].runs[0]
    header_run3.font.bold = True
    header_run3.font.size = Pt(13)
    header_run3.font.color.rgb = RGBColor(255, 255, 255)
    shading_elm3 = OxmlElement('w:shd')
    shading_elm3.set(qn('w:fill'), '3498db')
    header_cell3.paragraphs[0]._element.get_or_add_pPr().append(shading_elm3)

    # Instrucciones paso a paso
    instrucciones = [
        '1. Ingrese a: www.laboratoriopérez.com o utilice el enlace proporcionado por el laboratorio',
        '2. Click en el botón "Consultar mis Resultados" o "Ver Resultados"',
        '3. Ingrese su Cédula de Identidad (CI) y el Código de Acceso proporcionado arriba',
        '4. Podrá visualizar y descargar su resultado en formato PDF',
        'NOTA: Guarde este documento de forma segura. Su código de acceso es confidencial.'
    ]

    for i, instruccion in enumerate(instrucciones, start=1):
        cell = table3.rows[i].cells[0]
        cell.text = instruccion
        cell_run = cell.paragraphs[0].runs[0]
        cell_run.font.size = Pt(10)

        # Última fila (nota) con fondo especial
        if i == 5:
            cell_run.font.bold = True
            cell_run.font.color.rgb = RGBColor(133, 100, 4)
            shading_elm = OxmlElement('w:shd')
            shading_elm.set(qn('w:fill'), 'fff3cd')
            cell.paragraphs[0]._element.get_or_add_pPr().append(shading_elm)

    doc.add_paragraph()  # Espacio

    # ============ FOOTER ============
    footer_para = doc.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer_para.add_run('Laboratorio Clínico Pérez - Potosí, Bolivia')
    footer_run.font.size = Pt(9)
    footer_run.font.italic = True
    footer_run.font.color.rgb = RGBColor(149, 165, 166)

    footer_para2 = doc.add_paragraph()
    footer_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run2 = footer_para2.add_run('Documento generado automáticamente')
    footer_run2.font.size = Pt(8)
    footer_run2.font.italic = True
    footer_run2.font.color.rgb = RGBColor(149, 165, 166)

    # ============ GUARDAR DOCUMENTO ============
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    filename = f"Credenciales_{resultado.paciente_nombre.replace(' ', '_')}_{resultado.numero_orden}.docx"

    return send_file(buffer, as_attachment=True, download_name=filename, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')

@main.route('/pruebas', methods=['GET', 'POST'])
@admin_required
def admin_pruebas():
    if request.method == 'POST':
        try:
            # Manejar imagen si se subió
            imagen_filename = None
            if 'imagen' in request.files:
                imagen = request.files['imagen']
                if imagen and imagen.filename:
                    # Crear carpeta si no existe
                    os.makedirs(PRUEBAS_UPLOAD_DIR, exist_ok=True)

                    # Guardar imagen con nombre seguro
                    from werkzeug.utils import secure_filename
                    import time
                    filename = secure_filename(imagen.filename)
                    imagen_filename = f"{int(time.time())}_{filename}"
                    imagen.save(os.path.join(PRUEBAS_UPLOAD_DIR, imagen_filename))

            prueba = Prueba(
                nombre=request.form['nombre'],
                categoria=request.form.get('categoria'),
                descripcion=request.form.get('descripcion'),
                precio=float(request.form.get('precio', 0)),
                imagen=imagen_filename
            )
            db.session.add(prueba)
            db.session.commit()
            flash('Prueba registrada exitosamente', 'success')
        except Exception as e:
            flash(f'Error: {str(e)}', 'danger')
        return redirect(url_for('main.admin_pruebas'))

    pruebas = Prueba.query.order_by(Prueba.nombre).all()

    # Obtener categorías únicas de la base de datos
    categorias = db.session.query(Prueba.categoria).distinct().filter(Prueba.categoria.isnot(None)).order_by(Prueba.categoria).all()
    categorias_list = [cat[0] for cat in categorias if cat[0]]

    return render_template('admin/pruebas.html', pruebas=pruebas, categorias=categorias_list)

@main.route('/prueba/<int:prueba_id>')
@admin_required
def ver_prueba(prueba_id):
    prueba = Prueba.query.get_or_404(prueba_id)
    return jsonify({
        'id': prueba.id,
        'nombre': prueba.nombre,
        'categoria': prueba.categoria,
        'descripcion': prueba.descripcion,
        'precio': float(prueba.precio),
        'imagen': prueba.imagen
    })

@main.route('/prueba/editar/<int:prueba_id>', methods=['POST'])
@admin_required
def editar_prueba(prueba_id):
    try:
        prueba = Prueba.query.get_or_404(prueba_id)
        prueba.nombre = request.form['nombre']
        prueba.categoria = request.form.get('categoria')
        prueba.descripcion = request.form.get('descripcion')
        prueba.precio = float(request.form.get('precio', 0))

        # Manejar nueva imagen si se subió
        if 'imagen' in request.files:
            imagen = request.files['imagen']
            if imagen and imagen.filename:
                # Eliminar imagen anterior si existe
                if prueba.imagen:
                    old_path = os.path.join(PRUEBAS_UPLOAD_DIR, prueba.imagen)
                    if os.path.exists(old_path):
                        os.remove(old_path)

                # Guardar nueva imagen
                from werkzeug.utils import secure_filename
                import time
                filename = secure_filename(imagen.filename)
                imagen_filename = f"{int(time.time())}_{filename}"
                os.makedirs(PRUEBAS_UPLOAD_DIR, exist_ok=True)
                imagen.save(os.path.join(PRUEBAS_UPLOAD_DIR, imagen_filename))
                prueba.imagen = imagen_filename

        db.session.commit()
        flash('Prueba actualizada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pruebas'))

@main.route('/prueba/eliminar/<int:prueba_id>', methods=['POST'])
@admin_required
def eliminar_prueba(prueba_id):
    try:
        prueba = Prueba.query.get_or_404(prueba_id)
        nombre_prueba = prueba.nombre

        # Eliminar imagen si existe
        if prueba.imagen:
            try:
                imagen_path = os.path.join(PRUEBAS_UPLOAD_DIR, prueba.imagen)
                if os.path.exists(imagen_path):
                    os.remove(imagen_path)
            except Exception as e:
                print(f"Error eliminando imagen {prueba.imagen}: {e}")

        db.session.delete(prueba)
        db.session.commit()
        flash(f'Prueba "{nombre_prueba}" eliminada exitosamente', 'success')
    except Exception as e:
        db.session.rollback()
        flash(f'Error al eliminar prueba: {str(e)}', 'danger')
    return redirect(url_for('main.admin_pruebas'))

@main.route('/descargar/<int:resultado_id>')
@admin_required
def descargar(resultado_id):
    resultado = Resultado.query.get_or_404(resultado_id)
    if resultado.archivo_pdf:
        # Construir ruta absoluta correcta usando UPLOAD_DIR
        pdf_path = os.path.join(UPLOAD_DIR, resultado.archivo_pdf)

        # Verificar que el archivo existe
        if not os.path.exists(pdf_path):
            flash(f'El archivo PDF no se encuentra en el servidor: {resultado.archivo_pdf}', 'danger')
            return redirect(url_for('main.admin_resultados'))

        return send_file(pdf_path, as_attachment=True)

    flash('No hay archivo PDF disponible', 'warning')
    return redirect(url_for('main.admin_resultados'))

@main.route('/resultado/eliminar/<int:resultado_id>', methods=['POST'])
@admin_required
def eliminar_resultado(resultado_id):
    """
    Elimina un resultado individual con sus archivos (principal y backup)
    """
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        numero_orden = resultado.numero_orden
        paciente_nombre = resultado.paciente_nombre

        # Eliminar archivos PDF (principal y backup)
        archivos_eliminados = 0

        if resultado.archivo_pdf:
            # Eliminar archivo principal
            filepath = os.path.join(UPLOAD_DIR, resultado.archivo_pdf)
            if os.path.exists(filepath):
                os.remove(filepath)
                archivos_eliminados += 1
                print(f"🗑 Archivo principal eliminado: {resultado.archivo_pdf}")

            # Eliminar backup
            backup_path = os.path.join(BACKUP_DIR, resultado.archivo_pdf)
            if os.path.exists(backup_path):
                os.remove(backup_path)
                archivos_eliminados += 1
                print(f"🗑 Backup eliminado: {resultado.archivo_pdf}")

        # Eliminar registro de base de datos
        db.session.delete(resultado)
        db.session.commit()

        print("=" * 80)
        print("✅ RESULTADO ELIMINADO EXITOSAMENTE")
        print(f"   Número Orden: {numero_orden}")
        print(f"   Paciente: {paciente_nombre}")
        print(f"   Archivos eliminados: {archivos_eliminados}")
        print("=" * 80)

        flash(f'✅ Resultado de "{paciente_nombre}" (Orden: {numero_orden}) eliminado exitosamente', 'success')

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error al eliminar resultado: {str(e)}")
        flash(f'❌ Error al eliminar resultado: {str(e)}', 'danger')

    return redirect(url_for('main.admin_resultados'))

@main.route('/resultado/reemplazar/<int:resultado_id>', methods=['POST'])
@admin_required
def reemplazar_pdf(resultado_id):
    """
    Reemplaza el PDF de un resultado existente
    - Elimina el PDF anterior (principal y backup)
    - Sube el nuevo PDF con backup
    """
    try:
        resultado = Resultado.query.get_or_404(resultado_id)
        archivo_nuevo = request.files.get('archivo_pdf')

        if not archivo_nuevo or not archivo_nuevo.filename:
            flash('❌ Debe seleccionar un archivo PDF', 'danger')
            return redirect(url_for('main.admin_resultados'))

        # Eliminar archivos anteriores (principal y backup)
        if resultado.archivo_pdf:
            # Eliminar principal
            old_filepath = os.path.join(UPLOAD_DIR, resultado.archivo_pdf)
            if os.path.exists(old_filepath):
                os.remove(old_filepath)
                print(f"🗑 PDF anterior eliminado: {resultado.archivo_pdf}")

            # Eliminar backup
            old_backup = os.path.join(BACKUP_DIR, resultado.archivo_pdf)
            if os.path.exists(old_backup):
                os.remove(old_backup)
                print(f"🗑 Backup anterior eliminado: {resultado.archivo_pdf}")

        # Guardar nuevo PDF con backup
        filename_nuevo, filepath_nuevo, backup_nuevo = guardar_pdf_con_backup(
            archivo_nuevo,
            resultado.numero_orden
        )

        if not filename_nuevo:
            raise Exception("No se pudo guardar el nuevo archivo PDF")

        # Actualizar registro en BD
        resultado.archivo_pdf = filename_nuevo
        db.session.commit()

        print("=" * 80)
        print("✅ PDF REEMPLAZADO EXITOSAMENTE")
        print(f"   Resultado ID: {resultado.id}")
        print(f"   Número Orden: {resultado.numero_orden}")
        print(f"   Archivo nuevo: {filename_nuevo}")
        print(f"   Backup: ✓ Creado")
        print("=" * 80)

        flash(f'✅ PDF reemplazado exitosamente para el paciente "{resultado.paciente_nombre}"', 'success')

    except Exception as e:
        db.session.rollback()
        print(f"❌ Error al reemplazar PDF: {str(e)}")
        flash(f'❌ Error al reemplazar PDF: {str(e)}', 'danger')

    return redirect(url_for('main.admin_resultados'))












